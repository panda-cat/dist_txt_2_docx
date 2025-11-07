import os
import glob
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from ntc_templates.parse import parse_output
import argparse
import sys

# Set NET_TEXTFSM for frozen applications (PyInstaller)
if getattr(sys, 'frozen', False):
    # In PyInstaller, sys._MEIPASS is the temp dir for bundled files
    template_dir = os.path.join(sys._MEIPASS, 'ntc_templates', 'templates')
    os.environ['NET_TEXTFSM'] = template_dir

# --- 平台检测函数 (已根据您的反馈修正) ---
def detect_platform(raw_text):
    """
    通过文本中的关键字检测平台
    """
    text_lower = raw_text.lower()
   
    if 'cisco' in text_lower:
        return 'cisco_ios'
   
    # 感谢您的指正：ntc_templates 使用 'hp_comware' 作为 H3C 的 key
    # 'display' 是 H3C 的命令, 'h3c' 是品牌
    elif 'h3c' in text_lower or 'display' in text_lower:
        return 'hp_comware' # <-- 已修正
   
    else:
        # 如果没有明确线索，尝试根据 show/display 数量猜测
        if text_lower.count('show ') > text_lower.count('display '):
             return 'cisco_ios'
        return None

# --- 关键修正：按“提示符”分割命令 ---
# 匹配提示符的正则表达式：
# (?m) : 多行模式，^ 匹配行首
# ^ : 行首
# ([\w\d\.\-]+[#>] ?.*) : 捕获组 1 (提示符+命令)
# [\w\d\.\-]+ : 匹配主机名 (如 'My-Switch.Cisco' 或 'H3C-Core')
# [#>] : 匹配提示符结尾 '#' 或 '>'
# ?.* : 匹配提示符后的空格和整个命令 (如 ' show version')
PROMPT_REGEX = re.compile(r'(?m)^([\w\d\.\-]+[#>] ?.*)')

# 用于从 "Switch# show version" 中提取 "show version"
COMMAND_LINE_REGEX = re.compile(r'^[\w\d\.\-]+[#>] ?')

def split_commands(raw_text):
    """
    使用正则表达式按设备提示符将原始文本分割为 (命令, 输出) 元组
    """
    commands = []
   
    # 使用 .split()，我们将得到 [登录信息, 提示符+命令1, 输出1, 提示符+命令2, 输出2, ...]
    sections = PROMPT_REGEX.split(raw_text)
   
    if len(sections) < 3:
        # 如果没有匹配到任何提示符
        return []
   
    # 登录信息/Banner 在 sections[0]，我们丢弃它
    # 我们从索引 1 开始，步长为 2
    for i in range(1, len(sections), 2):
        cmd_line = sections[i].strip()
       
        # 提取真正的命令 (移除提示符)
        cmd = COMMAND_LINE_REGEX.sub('', cmd_line).strip()
       
        # 确保我们没有抓到空命令 (例如用户只敲了回车)
        if not cmd:
            continue
           
        output = sections[i+1].strip()
        commands.append((cmd, output))
       
    return commands

# --- 主逻辑 ---
def generate_report(txt_dir='devices', output_docx='network_report.docx'):
    doc = Document()
   
    # 设置一个基础字体
    style = doc.styles['Normal']
    style.font.name = 'Calibri' # 或者 '微软雅黑'
    style.font.size = Pt(10.5)
    doc.add_heading('网络设备报告', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
   
    txt_files = glob.glob(os.path.join(txt_dir, '*.txt'))
    if not txt_files:
        print(f"在 '{txt_dir}' 目录中未找到 .txt 文件。")
        return
    for txt_file in txt_files:
        device_name = os.path.basename(txt_file).replace('.txt', '')
        print(f"\n--- 正在处理设备: {device_name} ---")
       
        try:
            with open(txt_file, 'r', encoding='utf-8') as f:
                raw_text = f.read()
        except Exception as e:
            print(f" [!] 读取文件 {txt_file} 失败: {e}")
            doc.add_paragraph(f"设备 {device_name}：读取文件失败: {e}")
            continue
       
        platform = detect_platform(raw_text)
        if not platform:
            print(f" [!] 设备 {device_name}：无法检测平台，跳过。")
            doc.add_paragraph(f"设备 {device_name}：无法检测平台，跳过。")
            continue
       
        doc.add_heading(f"设备：{device_name} (平台: {platform})", level=2)
       
        commands = split_commands(raw_text)
        if not commands:
            print(f" [!] 在 {device_name} 中未分割出任何命令。")
            doc.add_paragraph("未找到任何命令输出。")
            continue
        for cmd, output in commands:
            print(f" [+] 正在解析: {cmd}")
            doc.add_heading(f"命令: {cmd}", level=3)
           
            try:
                # ntc_templates 会自动根据 platform 和 command 字符串
                # 找到 'cisco_ios_show_version.template' 或 'hp_comware_display_cpu.template'
                parsed_data = parse_output(platform=platform, command=cmd, data=output)
               
                if not parsed_data:
                    doc.add_paragraph("未解析到数据 (ntc_templates 未返回结果，可能是不支持的命令或无匹配)。")
                    print(" ... 解析无数据。")
                    continue
               
                # 'parsed_data' 是一个字典列表 [ {dict1}, {dict2}, ... ]
                headers = list(parsed_data[0].keys())
               
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = True
               
                # 表头
                hdr_cells = table.rows[0].cells
                for i, key in enumerate(headers):
                    hdr_cells[i].text = key.upper()
                    hdr_cells[i].paragraphs[0].runs[0].bold = True
               
                # 数据行
                for row_data in parsed_data:
                    row_cells = table.add_row().cells
                    for i, key in enumerate(headers):
                        row_cells[i].text = str(row_data.get(key, ''))
               
                doc.add_paragraph() # 表格后添加空行
                       
            except Exception as e:
                print(f" [!] 解析或制表失败: {e}")
                doc.add_paragraph(f"解析或生成表格时出错：{str(e)}")
       
        doc.add_page_break() # 每个设备后分页
   
    # 保存 DOCX
    try:
        # 获取输出目录
        output_dir = os.path.dirname(output_docx)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
        doc.save(output_docx)
        print(f"\n=========================================")
        print(f"✅ 报告已成功生成: {output_docx}")
        print(f"=========================================")
    except PermissionError:
        print(f"\n[X] 错误: 保存失败！请关闭已打开的 '{output_docx}' 文件。")
    except Exception as e:
        print(f"\n[X] 错误: 保存报告失败: {e}")

# 运行脚本
if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='生成网络设备报告')
    parser.add_argument('--input', '-i', type=str, default='devices', help='TXT文件所在的文件夹 (默认: devices)')
    parser.add_argument('--output', '-o', type=str, default='network_report.docx', help='输出DOCX文件路径 (默认: network_report.docx)')
    args = parser.parse_args()

    # 确保输入目录存在
    if not os.path.exists(args.input):
        os.makedirs(args.input)
        print(f"创建 '{args.input}' 目录，请将 .txt 文件放入其中。")
    else:
        generate_report(txt_dir=args.input, output_docx=args.output)
