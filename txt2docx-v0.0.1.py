# -*- coding: utf-8 -*-
"""
批量读取 *.txt（IP 命名），按 IP 顺序生成 1 个 Word 文档
每台交换机一页，支持：
1. 全新表格报告（无模板时）
2. 模板占位符替换（提供模板时）

命令行参数：
  -i, --input     存放 TXT 文件的文件夹，默认当前目录
  -t, --template  Word 模板文件路径（可选）
  -o, --output    输出 Word 文件路径（必填）

模板占位符示例：
IP 地址：{IP}
主机名：{HOSTNAME}
运行时间：{UPTIME}
序列号：{SN}
CPU 使用率：{CPU_UTILIZATION}
内存使用率：{MEMORY_UTILIZATION}
NTP 状态：{NTP_STATUS}
报告生成时间：{REPORT_TIME}
"""
import re
import argparse
import ipaddress
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn


# ------------------- 分页符 -------------------
def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


# ------------------- 占位符替换（保留格式） -------------------
def replace_in_run(run, placeholder_map):
    for ph, val in placeholder_map.items():
        run.text = run.text.replace(f"{{{ph}}}", str(val))


def replace_in_paragraph(para, placeholder_map):
    for run in para.runs:
        replace_in_run(run, placeholder_map)


def replace_in_cell(cell, placeholder_map):
    for paragraph in cell.paragraphs:
        replace_in_paragraph(paragraph, placeholder_map)


# ------------------- 解析单个 txt -------------------
def parse_switch(txt_path: Path) -> dict:
    """返回 {key: value}，key 为小写下划线形式"""
    data = {"_filename": txt_path.stem}   # 记录 IP
    with txt_path.open(encoding="utf-8", errors="ignore") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            m = re.match(r"^\s*(.+?)\s*:\s*(.+?)\s*$", line)
            if m:
                raw_key, raw_val = m.groups()
                key = (raw_key.lower()
                       .replace(" ", "_")
                       .replace("serial_number", "sn")
                       .replace("cpu_utilization", "cpu_utilization")
                       .replace("memory_utilization", "memory_utilization"))
                data[key] = raw_val.strip()
    return data


# ------------------- 按 IP 排序 -------------------
def sort_by_ip(file_list):
    """返回按 IP 地址自然排序的文件列表"""
    def ip_key(p):
        try:
            return ipaddress.ip_address(p.stem)
        except ValueError:
            return ipaddress.ip_address("255.255.255.255")  # 非法 IP 放最后
    return sorted(file_list, key=ip_key)


# ------------------- 生成全新 Word（每页一台） -------------------
def generate_multi_word(switches, out_path: Path):
    doc = Document()
    # 设置中文字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    report_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    for i, data in enumerate(switches):
        if i > 0:
            add_page_break(doc)

        ip = data.get("_filename", "Unknown")
        hostname = data.get("hostname", "N/A")

        # 标题
        title = doc.add_heading(f"交换机状态报告 - {ip} ({hostname})", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 表格
        table = doc.add_table(rows=1, cols=2, style="Table Grid")
        hdr = table.rows[0].cells
        hdr[0].text = "项目"
        hdr[1].text = "值"
        for cell in hdr:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.size = Pt(11)

        order = [
            ("hostname", "主机名"),
            ("uptime", "运行时间"),
            ("sn", "序列号"),
            ("cpu_utilization", "CPU 使用率"),
            ("memory_utilization", "内存使用率"),
            ("ntp_status", "NTP 状态"),
        ]

        for key, label in order:
            if key in data:
                row = table.add_row().cells
                row[0].text = label
                row[1].text = data[key]

        # 底部时间
        doc.add_paragraph(f"报告生成时间：{report_time}")

    doc.save(out_path)
    print(f"全新 Word 已生成（每页一台）：{out_path}")


# ------------------- 模板替换（每页插入模板） -------------------
def replace_template_multi(switches, tmpl_path: Path, out_path: Path):
    template_doc = Document(tmpl_path)
    num_template_paras = len(template_doc.paragraphs)
    num_template_tables = len(template_doc.tables)

    doc = Document()
    # 设置中文字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    report_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    for i, data in enumerate(switches):
        if i > 0:
            add_page_break(doc)

        # 复制模板内容
        for element in template_doc.element.body:
            doc.element.body.append(element)

        # 构造占位符映射（大写）
        placeholder_map = {k.upper(): v for k, v in data.items() if not k.startswith("_")}
        placeholder_map["IP"] = data["_filename"]
        placeholder_map["REPORT_TIME"] = report_time

        # 替换当前页的段落（保留格式）
        new_paras = doc.paragraphs[-num_template_paras:]
        for para in new_paras:
            replace_in_paragraph(para, placeholder_map)

        # 替换当前页的表格（保留格式）
        new_tables = doc.tables[-num_template_tables:]
        for table in new_tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_cell(cell, placeholder_map)

    doc.save(out_path)
    print(f"模板替换完成（每页一台）：{out_path}")


# ------------------- 主程序 -------------------
def main():
    parser = argparse.ArgumentParser(
        description="""批量读取 *.txt（IP 命名），按 IP 顺序生成 1 个 Word 文档
每台交换机一页，支持：
1. 全新表格报告（无模板时）
2. 模板占位符替换（提供模板时）""",
        formatter_class=argparse.RawDescriptionHelpFormatter
    )
    parser.add_argument('-i', '--input', default='.', help='存放 TXT 文件的文件夹，默认当前目录')
    parser.add_argument('-t', '--template', help='Word 模板文件路径（可选，若提供则使用模板替换模式）')
    parser.add_argument('-o', '--output', required=True, help='输出 Word 文件路径')

    args = parser.parse_args()

    txt_dir = Path(args.input)
    if not txt_dir.exists() or not txt_dir.is_dir():
        print(f"输入目录不存在或不是目录：{txt_dir}")
        return

    pattern = re.compile(r"^\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\.txt$")
    txt_files = [p for p in txt_dir.iterdir() if p.is_file() and pattern.match(p.name)]
    if not txt_files:
        print("未找到任何 *.txt 文件（IP 命名）")
        return

    txt_files = sort_by_ip(txt_files)
    print(f"找到 {len(txt_files)} 个文件，按 IP 排序：")
    for f in txt_files:
        print(f"  → {f.name}")

    switches = [parse_switch(txt_file) for txt_file in txt_files]

    out_path = Path(args.output)
    if args.template:
        tmpl_path = Path(args.template)
        if not tmpl_path.exists():
            print(f"模板不存在：{tmpl_path}")
            return
        replace_template_multi(switches, tmpl_path, out_path)
    else:
        generate_multi_word(switches, out_path)


if __name__ == "__main__":
    main()
