# -*- coding: utf-8 -*-
"""
批量读取 *.txt 文件（以IP地址开头命名），按 IP 顺序生成 1 个 Word 文档。
每台交换机一页，支持：
1. 全新表格报告（无模板时）
2. 模板占位符替换（提供模板时）

[!] 核心功能:
- 支持灵活的文件命名（例如 `192.168.1.1.txt`, `192.168.1.1-backup.txt`）。
- 自动识别并解析 Cisco (IOS/XE), Huawei (VRP), H3C (Comware) 设备。
- 适配 Cisco StackWise, Huawei iStack, H3C IRF 堆叠/虚拟化技术。
- 采用优先级策略精确解析主机名，有效避免 LLDP 等信息干扰。
- 支持在 Word 模板中使用特殊占位符 {MEMBER_TABLE} 动态生成堆叠成员表格。

命令行参数：
  -i, --input     存放 TXT 文件的文件夹，默认当前目录
  -t, --template  Word 模板文件路径（可选）
  -o, --output    输出 Word 文件路径（必填）
"""
import re
import argparse
import ipaddress
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from copy import deepcopy

# ------------------- Word 文档操作辅助函数 -------------------

def add_page_break(doc):
    """添加分页符"""
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def set_default_font(doc, font_name='微软雅黑'):
    """设置文档默认中英文字体"""
    style = doc.styles['Normal']
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def replace_text_in_doc(doc, placeholder_map):
    """在整个文档的段落和表格中替换占位符（保留格式）"""
    for para in doc.paragraphs:
        for run in para.runs:
            for ph, val in placeholder_map.items():
                run.text = run.text.replace(f"{{{ph}}}", str(val))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for ph, val in placeholder_map.items():
                            run.text = run.text.replace(f"{{{ph}}}", str(val))

# ------------------- 核心解析逻辑：各厂商解析器 -------------------

def parse_cisco(content: str) -> dict:
    """解析 Cisco IOS/IOS-XE 设备信息"""
    data = {"vendor": "Cisco", "members": []}
    
    # --- 主机名解析（采用优先级策略，抗LLDP干扰） ---
    hostname = None
    # 策略1: 优先从命令提示符中寻找主机名 (e.g., Switch#show...)
    m_prompt = re.search(r"(\S+?)#", content, re.I)
    if m_prompt:
        hostname = m_prompt.group(1).strip()
    
    # 策略2: 如果策略1失败, 则查找以 'hostname' 开头的配置行
    if not hostname:
        m_config = re.search(r"^\s*hostname\s+(.+?)\s*$", content, re.I | re.M)
        if m_config:
            hostname = m_config.group(1).strip()
            
    data["hostname"] = hostname or "N/A"
    
    # --- 公共信息提取 ---
    m = re.search(r"uptime is (.+)", content, re.I)
    data["uptime"] = m.group(1).strip() if m else "N/A"
    m = re.search(r"Clock is (.+)", content, re.I)
    data["ntp_status"] = m.group(1).strip() if m else "N/A"
    m_cpu = re.search(r"CPU utilization for five seconds: (\S+)", content, re.I)
    data["cpu_utilization"] = m_cpu.group(1).split('/')[0] if m_cpu else "N/A"
    m_total = re.search(r"Processor Pool Total:\s+(\d+)", content, re.I)
    m_used = re.search(r"Used:\s+(\d+)", content, re.I)
    if m_total and m_used:
        total, used = int(m_total.group(1)), int(m_used.group(1))
        data["memory_utilization"] = f"{(used / total * 100):.2f}%" if total > 0 else "0%"
    else:
        data["memory_utilization"] = "N/A"

    # --- 堆叠/成员信息 (关键) ---
    if 'show switch' in content.lower():
        member_matches = re.finditer(
            r"^\s*([*\d])\s+\S+\s+([\w-]+)\s+([0-9a-f:.]+)\s+(\w+)", content, re.M | re.I)
        members = [
            {"id": m.group(1).replace('*','').strip(), "model": m.group(2), "mac_address": m.group(3), "status": m.group(4), "sn": "N/A"}
            for m in member_matches
        ]
        sn_matches = re.finditer(r"Switch\s+(\d+)\s+SERIAL NUMBER\s+:\s+(\S+)", content, re.I | re.M)
        sn_map = {m.group(1): m.group(2) for m in sn_matches}
        for member in members:
            member["sn"] = sn_map.get(member["id"], "N/A")
        
        if members:
            data["members"] = members
            data["is_stack"] = len(members) > 1
    
    # --- 单台设备信息 (作为成员=1的特例处理) ---
    if not data["members"]:
        member = {"id": "1", "status": "Ready", "cpu": data.get("cpu_utilization", "N/A"), "memory": data.get("memory_utilization", "N/A")}
        m = re.search(r"System Serial Number\s+:\s+(\S+)", content, re.I)
        member["sn"] = m.group(1) if m else "N/A"
        data["sn"] = member["sn"]
        m = re.search(r"Model Number\s+:\s+(\S+)", content, re.I)
        member["model"] = m.group(1) if m else "N/A"
        data["model"] = member["model"]
        m = re.search(r"Cisco IOS.*, Version (\S+),", content, re.I)
        data["ios_version"] = m.group(1) if m else "N/A"
        data["members"].append(member)
        data["is_stack"] = False
        
    return data

def parse_huawei(content: str) -> dict:
    """解析 Huawei VRP 设备信息"""
    data = {"vendor": "Huawei", "members": []}

    # --- 主机名解析（采用优先级策略，抗LLDP干扰） ---
    hostname = None
    m_prompt = re.search(r"<(\S+?)>", content, re.I)
    if m_prompt:
        hostname = m_prompt.group(1).strip()
    if not hostname:
        m_config = re.search(r"^\s*(?:sysname|hostname)\s+(.+?)\s*$", content, re.I | re.M)
        if m_config:
            hostname = m_config.group(1).strip()
    data["hostname"] = hostname or "N/A"

    # --- 公共信息提取 ---
    m = re.search(r"uptime is (.+)", content, re.I)
    data["uptime"] = m.group(1).strip() if m else "N/A"
    m = re.search(r"clock status\s*:\s*(.+)", content, re.I)
    data["ntp_status"] = m.group(1).strip() if m else "N/A"
    m = re.search(r"Version \d\.\d+ \((.+?)\)", content, re.I)
    data["ios_version"] = m.group(1) if m else "N/A"
    
    # --- 堆叠/成员信息 ---
    if 'display device' in content.lower():
        member_matches = re.finditer(
            r"^\s*(\d+)\s+(\w+)\s+\w+\s+([\w-]+)\s+([0-9A-Z]+)", content, re.M | re.I)
        members = [
            {"id": m.group(1), "role": m.group(2), "model": m.group(3), "sn": m.group(4), "cpu": "N/A", "memory": "N/A"}
            for m in member_matches
        ]
        cpu_matches = re.finditer(r"CPU Usage for Slot\s+(\d+)\s+is\s+(\S+)", content, re.I | re.M)
        cpu_map = {m.group(1): m.group(2) for m in cpu_matches}
        mem_matches = re.finditer(r"Memory usage of slot\s+(\d+):\s+(\S+)", content, re.I | re.M)
        mem_map = {m.group(1): m.group(2) for m in mem_matches}
        for member in members:
            member["cpu"] = cpu_map.get(member["id"], "N/A")
            member["memory"] = mem_map.get(member["id"], "N/A")
            
        if members:
            data["members"] = members
            data["is_stack"] = len(members) > 1
            for member in members:
                if member.get('role', '').lower() == 'master':
                    data['cpu_utilization'] = member.get('cpu', 'N/A')
                    data['memory_utilization'] = member.get('memory', 'N/A')
                    data['sn'] = member.get('sn', 'N/A')
                    data['model'] = member.get('model', 'N/A')
                    break
    
    # --- 单台设备信息 (作为成员=1的特例处理) ---
    if not data["members"]:
        member = {"id": "1"}
        m = re.search(r"BARCODE\s+:\s+(\S+)", content, re.I | re.M)
        member['sn'] = m.group(1) if m else "N/A"
        m = re.search(r"ITEM\s+:\s+(\S+)", content, re.I | re.M)
        member['model'] = m.group(1) if m else "N/A"
        m_cpu = re.search(r"Control Plane\s+CPU Usage is\s*(\S+)", content)
        member['cpu'] = m_cpu.group(1) if m_cpu else 'N/A'
        m_mem = re.search(r"Memory Using Percentage Is\s*(\S+)", content)
        member['memory'] = m_mem.group(1) if m_mem else 'N/A'
        data.update({'sn': member['sn'], 'model': member['model'], 'cpu_utilization': member['cpu'], 'memory_utilization': member['memory']})
        data["members"].append(member)
        data["is_stack"] = False

    return data

def parse_h3c(content: str) -> dict:
    """解析 H3C/HPE Comware 设备信息"""
    data = {"vendor": "H3C", "members": []}
    
    # --- 主机名解析（采用优先级策略，抗LLDP干扰） ---
    hostname = None
    m_prompt = re.search(r"<(\S+?)>", content, re.I)
    if m_prompt:
        hostname = m_prompt.group(1).strip()
    if not hostname:
        m_config = re.search(r"^\s*(?:sysname|hostname)\s+(.+?)\s*$", content, re.I | re.M)
        if m_config:
            hostname = m_config.group(1).strip()
    data["hostname"] = hostname or "N/A"

    # --- 公共信息提取 ---
    m = re.search(r"uptime is (.+)", content, re.I)
    data["uptime"] = m.group(1).strip() if m else "N/A"
    m = re.search(r"Clock status: (.+)", content, re.I)
    data["ntp_status"] = m.group(1).strip() if m else "N/A"
    m = re.search(r"Comware Software, Version ([\d\.]+)", content, re.I)
    data["ios_version"] = m.group(1) if m else "N/A"
    
    # --- 堆叠/成员信息 (IRF) ---
    if 'display irf' in content.lower() or 'display device' in content.lower():
        member_matches = re.finditer(r"^\s*(\d+)\s+(\w+)\s+([\w-]+)\s+([A-Z0-9]+)", content, re.M | re.I)
        members = [
            {"id": m.group(1), "role": m.group(2), "model": m.group(3), "sn": m.group(4), "cpu": "N/A", "memory": "N/A"}
            for m in member_matches
        ]
        cpu_matches = re.finditer(r"Slot\s+(\d+)\s+CPU usage:\s+(\S+)", content, re.I | re.M)
        cpu_map = {m.group(1): m.group(2) for m in cpu_matches}
        mem_matches = re.finditer(r"Slot\s+(\d+)\s+memory usage\s+\(Ratio\):\s+(\S+)", content, re.I | re.M)
        mem_map = {m.group(1): m.group(2) for m in mem_matches}
        for member in members:
            member["cpu"] = cpu_map.get(member["id"], "N/A")
            member["memory"] = mem_map.get(member["id"], "N/A")
            
        if members:
            data["members"] = members
            data["is_stack"] = len(members) > 1
            for member in members:
                if member.get('role', '').lower() == 'master':
                    data['cpu_utilization'] = member.get('cpu', 'N/A')
                    data['memory_utilization'] = member.get('memory', 'N/A')
                    data['sn'] = member.get('sn', 'N/A')
                    data['model'] = member.get('model', 'N/A')
                    break
    
    # --- 单台设备信息 (作为成员=1的特例处理) ---
    if not data["members"]:
        member = {"id": "1"}
        m_sn = re.search(r"Device serial number:\s*(\S+)", content, re.I)
        member['sn'] = m_sn.group(1) if m_sn else "N/A"
        m_model = re.search(r"Device model:\s*(\S+)", content, re.I)
        member['model'] = m_model.group(1) if m_model else "N/A"
        m_cpu = re.search(r"CPU average usage:\s*(\S+)", content, re.I)
        member['cpu'] = m_cpu.group(1) if m_cpu else "N/A"
        m_mem = re.search(r"Memory usage:\s*(\S+)", content, re.I)
        member['memory'] = m_mem.group(1) if m_mem else "N/A"
        data.update({'sn': member['sn'], 'model': member['model'], 'cpu_utilization': member['cpu'], 'memory_utilization': member['memory']})
        data["members"].append(member)
        data["is_stack"] = False

    return data

# ------------------- “控制中心”：自动检测并分发任务 -------------------

def parse_device_info(txt_path: Path, ip_address: str) -> dict:
    """自动检测设备类型并调用相应的解析器"""
    with txt_path.open(encoding="utf-8", errors="ignore") as f:
        content = f.read()

    data = {"_filename": ip_address}

    # --- 厂商“指纹”识别（使用多种独特关键词提高准确率） ---
    if re.search(r"Cisco IOS|\s#show", content, re.I):
        parsed_data = parse_cisco(content)
    elif re.search(r"VRP \(R\) software|HUAWEI|<HUAWEI>|display device", content, re.I):
        parsed_data = parse_huawei(content)
    elif re.search(r"Comware Software|<H3C>|display irf", content, re.I):
        parsed_data = parse_h3c(content)
    else:
        # Fallback: 如果无法识别厂商，作为 "Unknown" 并尝试基本键值对解析
        parsed_data = {"vendor": "Unknown", "is_stack": False, "members": [{}]}
        for line in content.splitlines():
            m = re.match(r"^\s*(.+?)\s*:\s*(.+?)\s*$", line.strip())
            if m:
                k, v = m.groups()
                # 兼容旧格式的hostname
                if k.lower() == 'hostname':
                    parsed_data['hostname'] = v.strip()
                parsed_data[k.lower().replace(" ", "_")] = v.strip()
    
    data.update(parsed_data)
    return data

# ------------------- 工具函数：按 IP 排序 -------------------

def sort_by_ip(file_list, ip_pattern):
    """根据从文件名中提取的IP地址进行自然排序"""
    def ip_key(p):
        match = ip_pattern.match(p.name)
        if match:
            ip_str = match.group(1)
            try:
                return ipaddress.ip_address(ip_str)
            except ValueError:
                return ipaddress.ip_address("255.255.255.255")
        return ipaddress.ip_address("255.255.255.255")
        
    return sorted(file_list, key=ip_key)

# ------------------- 报告生成模块 -------------------

def generate_multi_word(switches, out_path: Path):
    """生成全新的、包含多个设备页的Word报告"""
    doc = Document()
    set_default_font(doc)
    report_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    for i, data in enumerate(switches):
        if i > 0:
            add_page_break(doc)

        # 页面标题
        title = doc.add_heading(f"交换机状态报告 - {data.get('_filename', 'N/A')} ({data.get('hostname', 'N/A')})", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 表格 1: 设备概览
        doc.add_heading("设备概览", level=2)
        table_overview = doc.add_table(rows=2, cols=4, style="Table Grid")
        headers1 = ['主机名', '厂商', '主设备型号', '软件版本']
        values1 = [data.get(k, 'N/A') for k in ['hostname', 'vendor', 'model', 'ios_version']]
        for i, h in enumerate(headers1): table_overview.cell(0, i).text = h
        for i, v in enumerate(values1): table_overview.cell(1, i).text = str(v)

        # 表格 2: 运行状态
        table_status = doc.add_table(rows=2, cols=4, style="Table Grid")
        headers2 = ['运行时间', 'NTP状态', 'CPU使用率(主)', '内存使用率(主)']
        values2 = [data.get(k, 'N/A') for k in ['uptime', 'ntp_status', 'cpu_utilization', 'memory_utilization']]
        for i, h in enumerate(headers2): table_status.cell(0, i).text = h
        for i, v in enumerate(values2): table_status.cell(1, i).text = str(v)
        
        # 表格 3: 堆叠/IRF成员详情
        doc.add_heading("成员设备详情", level=2)
        members = data.get("members", [])
        if members:
            member_headers = ["ID/Slot", "角色", "型号", "序列号", "CPU", "内存", "状态"]
            table_members = doc.add_table(rows=1, cols=len(member_headers), style="Table Grid")
            for i, h in enumerate(member_headers): table_members.rows[0].cells[i].text = h
            for member in members:
                row_cells = table_members.add_row().cells
                row_cells[0].text = str(member.get("id", "N/A"))
                row_cells[1].text = member.get("role", "N/A")
                row_cells[2].text = member.get("model", "N/A")
                row_cells[3].text = member.get("sn", "N/A")
                row_cells[4].text = str(member.get("cpu", "N/A"))
                row_cells[5].text = str(member.get("memory", "N/A"))
                row_cells[6].text = member.get("status", "N/A")
        
        doc.add_paragraph(f"\n报告生成时间：{report_time}")

    doc.save(out_path)
    print(f"全新 Word 报告已生成：{out_path}")

def create_member_table_xml(doc, members_data):
    """为模板替换功能动态创建一个成员表格的XML对象"""
    if not members_data:
        return None
    
    headers = ["ID/Slot", "角色", "型号", "序列号", "CPU", "内存", "状态"]
    # 创建一个临时表格来构建结构和内容
    table = doc.add_table(rows=1, cols=len(headers), style="Table Grid")
    for i, h in enumerate(headers): table.rows[0].cells[i].text = h
    for member in members_data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(member.get("id", "N/A"))
        row_cells[1].text = member.get("role", "N/A")
        row_cells[2].text = member.get("model", "N/A")
        row_cells[3].text = member.get("sn", "N/A")
        row_cells[4].text = str(member.get("cpu", "N/A"))
        row_cells[5].text = str(member.get("memory", "N/A"))
        row_cells[6].text = member.get("status", "N/A")
        
    return table._tbl

def replace_template_multi(switches, tmpl_path: Path, out_path: Path):
    """使用模板生成报告，支持 {MEMBER_TABLE} 占位符"""
    doc = Document(tmpl_path)
    set_default_font(doc)
    template_body_elements = [deepcopy(element) for element in doc.element.body]

    # 清空原始文档，为生成多页内容做准备
    for element in list(doc.element.body):
        doc.element.body.remove(element)
    
    report_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    for i, data in enumerate(switches):
        if i > 0: add_page_break(doc)
        
        # 为当前设备复制一页模板内容
        for element in template_body_elements:
            doc.element.body.append(deepcopy(element))

        # 构造占位符映射表（大写）
        placeholder_map = {k.upper(): v for k, v in data.items() if not k.startswith("_") and not isinstance(v, (list, dict))}
        placeholder_map.update({
            "IP": data.get("_filename", "N/A"),
            "REPORT_TIME": report_time,
            "SN": data.get("sn", "N/A"),
            "CPU_UTILIZATION": data.get("cpu_utilization", "N/A"),
            "MEMORY_UTILIZATION": data.get("memory_utilization", "N/A"),
            "MODEL": data.get("model", "N/A"),
            "IOS_VERSION": data.get("ios_version", "N/A")
        })

        # 全局替换简单占位符
        replace_text_in_doc(doc, placeholder_map)

        # 特殊处理 {MEMBER_TABLE} 占位符
        for para in doc.paragraphs:
            if '{MEMBER_TABLE}' in para.text:
                para.text = "" # 清空占位符段落
                new_table_xml = create_member_table_xml(doc, data.get("members", []))
                if new_table_xml:
                    # 在占位符段落后插入新表格，然后删除占位符段落
                    p_element = para._p
                    p_element.addnext(new_table_xml)
                    parent = p_element.getparent()
                    parent.remove(p_element) # 从父节点移除表格（因为它已通过 addnext 附加）

    doc.save(out_path)
    print(f"模板替换完成：{out_path}")

# ------------------- 主程序入口 -------------------

def main():
    parser = argparse.ArgumentParser(
        description="批量读取TXT日志，生成多厂商、支持堆叠的Word巡检报告。",
        formatter_class=argparse.RawDescriptionHelpFormatter
    )
    parser.add_argument('-i', '--input', default='.', help='存放 TXT 文件的文件夹，默认当前目录')
    parser.add_argument('-t', '--template', help='Word 模板文件路径（可选，若提供则使用模板替换模式）')
    parser.add_argument('-o', '--output', required=True, help='输出 Word 文件路径')

    args = parser.parse_args()

    txt_dir = Path(args.input)
    if not txt_dir.is_dir():
        print(f"错误：输入路径不是一个有效的目录：{txt_dir}")
        return

    # 正则表达式，用于匹配“以IP地址开头”的txt文件，并捕获IP部分
    ip_pattern = re.compile(r"^(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}).*\.txt$")
    
    txt_files = [p for p in txt_dir.iterdir() if p.is_file() and ip_pattern.match(p.name)]
    if not txt_files:
        print(f"在目录 '{txt_dir}' 中未找到任何以IP地址开头的 .txt 文件。")
        return

    txt_files = sort_by_ip(txt_files, ip_pattern)
    print(f"找到 {len(txt_files)} 个文件，已按 IP 地址排序：")
    for f in txt_files:
        print(f"  → {f.name}")

    # 解析所有找到的文件
    switches = []
    for txt_file in txt_files:
        match = ip_pattern.match(txt_file.name)
        ip_address = match.group(1) if match else "0.0.0.0"
        switches.append(parse_device_info(txt_file, ip_address))

    # 根据是否提供模板，调用不同的报告生成函数
    out_path = Path(args.output)
    if args.template:
        tmpl_path = Path(args.template)
        if not tmpl_path.is_file():
            print(f"错误：模板文件不存在或不是一个有效文件：{tmpl_path}")
            return
        replace_template_multi(switches, tmpl_path, out_path)
    else:
        generate_multi_word(switches, out_path)

if __name__ == "__main__":
    main()
