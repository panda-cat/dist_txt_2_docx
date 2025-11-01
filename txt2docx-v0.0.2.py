# -*- coding: utf-8 -*-
"""
批量读取 *.txt（IP 命名），按 IP 顺序生成 1 个 Word 文档
每台交换机一页，支持：
1. 全新表格报告（无模板时）
2. 模板占位符替换（提供模板时）

[!] 升级版功能:
- 自动识别并支持 Cisco, Huawei, H3C 设备。
- 适配 Cisco StackWise, Huawei iStack, H3C IRF 堆叠交换机。
- 模板中支持 {MEMBER_TABLE} 占位符，用于自动生成堆叠成员信息表格。

命令行参数：
  -i, --input     存放 TXT 文件的文件夹，默认当前目录
  -t, --template  Word 模板文件路径（可选）
  -o, --output    输出 Word 文件路径（必填）

模板占位符示例：
IP 地址：{IP}
主机名：{HOSTNAME}
厂商：{VENDOR}
型号：{MODEL}
IOS版本：{IOS_VERSION}
运行时间：{UPTIME}
序列号：{SN}           (注: 对于堆叠设备，通常显示主设备序列号)
CPU 使用率：{CPU_UTILIZATION} (注: 对于堆叠设备，通常显示主设备CPU)
内存使用率：{MEMORY_UTILIZATION} (注: 对于堆叠设备，通常显示主设备内存)
NTP 状态：{NTP_STATUS}
报告生成时间：{REPORT_TIME}
堆叠成员表：{MEMBER_TABLE} (特殊占位符，将被替换为成员详情表格)
"""
import re
import argparse
import ipaddress
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from copy import deepcopy

# ------------------- Word 文档操作辅助函数 -------------------

def add_page_break(doc):
    """添加分页符"""
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def set_default_font(doc, font_name='微软雅黑'):
    """设置文档默认中英文字体"""
    style = doc.styles['Normal']
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def replace_text_in_doc(doc, placeholder_map):
    """在整个文档的段落和表格中替换占位符（保留格式）"""
    for para in doc.paragraphs:
        for run in para.runs:
            for ph, val in placeholder_map.items():
                run.text = run.text.replace(f"{{{ph}}}", str(val))
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for ph, val in placeholder_map.items():
                            run.text = run.text.replace(f"{{{ph}}}", str(val))


# ------------------- 核心解析逻辑 -------------------

def parse_cisco(content: str) -> dict:
    """解析 Cisco IOS/IOS-XE 设备信息"""
    data = {"vendor": "Cisco", "members": []}
    
    # --- 公共信息 ---
    # 主机名: Switch#
    m = re.search(r"(\S+?)#show", content, re.I)
    data["hostname"] = m.group(1) if m else "N/A"
    
    # 运行时间: Switch uptime is 1 year, 2 weeks, 3 days, 4 hours, 5 minutes
    m = re.search(r"uptime is (.+)", content, re.I)
    data["uptime"] = m.group(1).strip() if m else "N/A"
    
    # NTP: Clock is synchronized, stratum 3, reference is 10.0.0.1
    m = re.search(r"Clock is (.+)", content, re.I)
    data["ntp_status"] = m.group(1).strip() if m else "N/A"
    
    # 主CPU: CPU utilization for five seconds: 10%/1%; one minute: 9%; five minutes: 9%
    m = re.search(r"CPU utilization for five seconds: (\S+)", content, re.I)
    data["cpu_utilization"] = m.group(1).split('/')[0] if m else "N/A"
    
    # 主内存: Processor Pool Total: 12345678 Used: 1234567 Free: 1234567
    m_total = re.search(r"Processor Pool Total:\s+(\d+)", content, re.I)
    m_used = re.search(r"Used:\s+(\d+)", content, re.I)
    if m_total and m_used:
        total = int(m_total.group(1))
        used = int(m_used.group(1))
        data["memory_utilization"] = f"{(used / total * 100):.2f}%" if total > 0 else "0%"
    else:
        data["memory_utilization"] = "N/A"

    # --- 堆叠/成员信息 (关键) ---
    # `show switch` 命令是判断堆叠的关键
    if 'show switch' in content.lower():
        # 正则表达式匹配 show switch 的每一行
        #  1       Provision      WS-C3850-48P     00:11:22:33:44:55    Ready
        member_matches = re.finditer(
            r"^\s*([*\d])\s+\S+\s+([\w-]+)\s+([0-9a-f:.]+)\s+(\w+)", 
            content, re.M | re.I
        )
        members = []
        for match in member_matches:
            members.append({
                "id": match.group(1).replace('*','').strip(),
                "model": match.group(2),
                "mac_address": match.group(3),
                "status": match.group(4),
                "sn": "N/A" # 先占位
            })

        # 从 `show version` 中为每个成员找到序列号
        sn_matches = re.finditer(
            r"Switch\s+(\d+)\s+SERIAL NUMBER\s+:\s+(\S+)", 
            content, re.I | re.M
        )
        sn_map = {m.group(1): m.group(2) for m in sn_matches}
        
        for member in members:
            if member["id"] in sn_map:
                member["sn"] = sn_map[member["id"]]

        if members:
            data["members"] = members
            data["is_stack"] = len(members) > 1
    
    # 如果不是堆叠或没有 `show switch`，则作为单台设备处理
    if not data["members"]:
        member = {"id": "1", "status": "Ready", "cpu": data["cpu_utilization"], "memory": data["memory_utilization"]}
        m = re.search(r"System Serial Number\s+:\s+(\S+)", content, re.I)
        member["sn"] = m.group(1) if m else "N/A"
        data["sn"] = member["sn"] # 顶层SN
        
        m = re.search(r"Model Number\s+:\s+(\S+)", content, re.I)
        member["model"] = m.group(1) if m else "N/A"
        data["model"] = member["model"]
        
        m = re.search(r"Cisco IOS.*, Version (\S+),", content, re.I)
        data["ios_version"] = m.group(1) if m else "N/A"

        data["members"].append(member)
        data["is_stack"] = False

    return data


def parse_huawei(content: str) -> dict:
    """解析 Huawei VRP 设备信息"""
    data = {"vendor": "Huawei", "members": []}

    # --- 公共信息 ---
    # 主机名: <Switch>
    m = re.search(r"<(\S+?)>", content, re.I)
    data["hostname"] = m.group(1) if m else "N/A"
    
    # 运行时间: HUAWEI Uptime is 1 year, 2 weeks, 3 days, 4 hours, 5 minutes
    m = re.search(r"uptime is (.+)", content, re.I)
    data["uptime"] = m.group(1).strip() if m else "N/A"
    
    # NTP: clock status: synchronized
    m = re.search(r"clock status\s*:\s*(.+)", content, re.I)
    data["ntp_status"] = m.group(1).strip() if m else "N/A"
    
    # 版本: VRP (R) software, Version 5.170 (S6720 V200R011C10SPC600)
    m = re.search(r"Version \d\.\d+ \((.+?)\)", content, re.I)
    data["ios_version"] = m.group(1) if m else "N/A"
    
    # --- 堆叠/成员信息 ---
    # `display device` 是关键
    if 'display device' in content.lower():
        # 正则表达式匹配 display device 的每一行 (注意Slot Role Mac...的表头)
        # 1   Master  NORMAL    S5720-28P-LI-AC   1234567890ABCDEF   FAB
        member_matches = re.finditer(
            r"^\s*(\d+)\s+(\w+)\s+\w+\s+([\w-]+)\s+([0-9A-Z]+)", 
            content, re.M | re.I
        )
        members = []
        for match in member_matches:
             members.append({
                "id": match.group(1),
                "role": match.group(2),
                "model": match.group(3),
                "sn": match.group(4),
                "cpu": "N/A", # 先占位
                "memory": "N/A" # 先占位
            })

        # 解析每个 slot 的 CPU 和内存
        # display cpu-usage slot 1: 10%
        cpu_matches = re.finditer(r"CPU Usage for Slot\s+(\d+)\s+is\s+(\S+)", content, re.I | re.M)
        cpu_map = {m.group(1): m.group(2) for m in cpu_matches}
        
        # display memory-usage slot 1: 30%
        mem_matches = re.finditer(r"Memory usage of slot\s+(\d+):\s+(\S+)", content, re.I | re.M)
        mem_map = {m.group(1): m.group(2) for m in mem_matches}
        
        for member in members:
            member["cpu"] = cpu_map.get(member["id"], "N/A")
            member["memory"] = mem_map.get(member["id"], "N/A")

        if members:
            data["members"] = members
            data["is_stack"] = len(members) > 1
            # 将主设备信息提升到顶层
            for member in members:
                if member.get('role', '').lower() == 'master':
                    data['cpu_utilization'] = member.get('cpu', 'N/A')
                    data['memory_utilization'] = member.get('memory', 'N/A')
                    data['sn'] = member.get('sn', 'N/A')
                    data['model'] = member.get('model', 'N/A')
                    break
    
    # 单台设备 fallback
    if not data["members"]:
        member = {"id": "1"}
        m = re.search(r"DEVICE_NAME\s+:\s+(\S+)", content, re.I | re.M) # 通常在 esn.dat 中
        data['hostname'] = m.group(1) if m else data['hostname']
        
        m = re.search(r"BARCODE\s+:\s+(\S+)", content, re.I | re.M)
        member['sn'] = m.group(1) if m else "N/A"
        
        m = re.search(r"ITEM\s+:\s+(\S+)", content, re.I | re.M)
        member['model'] = m.group(1) if m else "N/A"

        m = re.search(r"Control Plane\s+CPU Usage is\s*(\S+)", content)
        member['cpu'] = m.group(1) if m else 'N/A'

        m = re.search(r"Memory Using Percentage Is\s*(\S+)", content)
        member['memory'] = m.group(1) if m else 'N/A'
        
        data.update({
            'sn': member['sn'],
            'model': member['model'],
            'cpu_utilization': member['cpu'],
            'memory_utilization': member['memory']
        })
        data["members"].append(member)
        data["is_stack"] = False

    return data


def parse_h3c(content: str) -> dict:
    """解析 H3C/HPE Comware 设备信息"""
    data = {"vendor": "H3C", "members": []}

    # --- 公共信息 ---
    m = re.search(r"<(\S+?)>", content, re.I)
    data["hostname"] = m.group(1) if m else "N/A"
    
    m = re.search(r"uptime is (.+)", content, re.I)
    data["uptime"] = m.group(1).strip() if m else "N/A"
    
    m = re.search(r"Clock status: (.+)", content, re.I)
    data["ntp_status"] = m.group(1).strip() if m else "N/A"
    
    # Comware Software, Version 7.1.064
    m = re.search(r"Comware Software, Version ([\d\.]+)", content, re.I)
    data["ios_version"] = m.group(1) if m else "N/A"

    # --- 堆叠/成员信息 (IRF) ---
    if 'display irf' in content.lower() or 'display device' in content.lower():
        # MemberID Role  Model                Serial Number
        # 1      Master S5130S-28S-EI         219801A0YNE19C000003
        member_matches = re.finditer(
            r"^\s*(\d+)\s+(\w+)\s+([\w-]+)\s+([A-Z0-9]+)", 
            content, re.M | re.I
        )
        members = []
        for match in member_matches:
             members.append({
                "id": match.group(1),
                "role": match.group(2),
                "model": match.group(3),
                "sn": match.group(4),
                "cpu": "N/A", 
                "memory": "N/A"
            })
        
        # Slot 1 CPU usage: 10%
        cpu_matches = re.finditer(r"Slot\s+(\d+)\s+CPU usage:\s+(\S+)", content, re.I | re.M)
        cpu_map = {m.group(1): m.group(2) for m in cpu_matches}
        
        # Slot 1 memory usage: 30%
        mem_matches = re.finditer(r"Slot\s+(\d+)\s+memory usage\s+\(Ratio\):\s+(\S+)", content, re.I | re.M)
        mem_map = {m.group(1): m.group(2) for m in mem_matches}
        
        for member in members:
            member["cpu"] = cpu_map.get(member["id"], "N/A")
            member["memory"] = mem_map.get(member["id"], "N/A")

        if members:
            data["members"] = members
            data["is_stack"] = len(members) > 1
            for member in members:
                if member.get('role', '').lower() == 'master':
                    data['cpu_utilization'] = member.get('cpu', 'N/A')
                    data['memory_utilization'] = member.get('memory', 'N/A')
                    data['sn'] = member.get('sn', 'N/A')
                    data['model'] = member.get('model', 'N/A')
                    break

    # 单台设备 fallback
    if not data["members"]:
        member = {"id": "1"}
        m = re.search(r"Device serial number:\s*(\S+)", content, re.I)
        member['sn'] = m.group(1) if m else "N/A"
        m = re.search(r"Device model:\s*(\S+)", content, re.I)
        member['model'] = m.group(1) if m else "N/A"

        m = re.search(r"CPU average usage:\s*(\S+)", content, re.I)
        member['cpu'] = m.group(1) if m else "N/A"
        m = re.search(r"Memory usage:\s*(\S+)", content, re.I)
        member['memory'] = m.group(1) if m else "N/A"

        data.update({
            'sn': member['sn'], 'model': member['model'],
            'cpu_utilization': member['cpu'], 'memory_utilization': member['memory']
        })
        data["members"].append(member)
        data["is_stack"] = False

    return data


def parse_device_info(txt_path: Path) -> dict:
    """自动检测设备类型并调用相应的解析器"""
    with txt_path.open(encoding="utf-8", errors="ignore") as f:
        content = f.read()

    data = {"_filename": txt_path.stem}  # 记录IP

    # Vendor detection
    if re.search(r"Cisco IOS Software|show version", content, re.I):
        parsed_data = parse_cisco(content)
    elif re.search(r"<HUAWEI>|display version", content, re.I):
        parsed_data = parse_huawei(content)
    elif re.search(r"<H3C>|Comware Software", content, re.I):
        parsed_data = parse_h3c(content)
    else:
        #  fallback to generic key-value parsing
        parsed_data = {"vendor": "Unknown", "is_stack": False, "members": [{}]}
        for line in content.splitlines():
            m = re.match(r"^\s*(.+?)\s*:\s*(.+?)\s*$", line.strip())
            if m:
                k, v = m.groups()
                parsed_data[k.lower().replace(" ", "_")] = v.strip()
    
    data.update(parsed_data)
    return data

# ------------------- 按 IP 排序 -------------------

def sort_by_ip(file_list):
    """返回按 IP 地址自然排序的文件列表"""
    def ip_key(p):
        try:
            return ipaddress.ip_address(p.stem)
        except ValueError:
            return ipaddress.ip_address("255.255.255.255") # 非法 IP 放最后
    return sorted(file_list, key=ip_key)


# ------------------- 生成全新 Word（每页一台） -------------------

def generate_multi_word(switches, out_path: Path):
    doc = Document()
    set_default_font(doc)
    report_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    for i, data in enumerate(switches):
        if i > 0:
            add_page_break(doc)

        ip = data.get("_filename", "N/A")
        hostname = data.get("hostname", "N/A")

        # 标题
        title = doc.add_heading(f"交换机状态报告 - {ip} ({hostname})", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # --- 表格 1: 公共信息 ---
        doc.add_heading("设备概览", level=2)
        table = doc.add_table(rows=1, cols=4, style="Table Grid")
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '主机名'
        hdr_cells[1].text = '厂商'
        hdr_cells[2].text = '主设备型号'
        hdr_cells[3].text = '软件版本'
        
        row_cells = table.add_row().cells
        row_cells[0].text = data.get('hostname', 'N/A')
        row_cells[1].text = data.get('vendor', 'N/A')
        row_cells[2].text = data.get('model', 'N/A')
        row_cells[3].text = data.get('ios_version', 'N/A')
        
        table = doc.add_table(rows=1, cols=4, style="Table Grid")
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '运行时间'
        hdr_cells[1].text = 'NTP状态'
        hdr_cells[2].text = 'CPU使用率(主)'
        hdr_cells[3].text = '内存使用率(主)'

        row_cells = table.add_row().cells
        row_cells[0].text = data.get('uptime', 'N/A')
        row_cells[1].text = data.get('ntp_status', 'N/A')
        row_cells[2].text = str(data.get('cpu_utilization', 'N/A'))
        row_cells[3].text = str(data.get('memory_utilization', 'N/A'))


        # --- 表格 2: 成员信息 ---
        doc.add_heading("成员设备详情", level=2)
        members = data.get("members", [])
        if members:
            member_headers = ["ID/Slot", "角色", "型号", "序列号", "CPU", "内存", "状态"]
            table = doc.add_table(rows=1, cols=len(member_headers), style="Table Grid")
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(member_headers):
                hdr_cells[i].text = header

            for member in members:
                row_cells = table.add_row().cells
                row_cells[0].text = str(member.get("id", "N/A"))
                row_cells[1].text = member.get("role", "N/A")
                row_cells[2].text = member.get("model", "N/A")
                row_cells[3].text = member.get("sn", "N/A")
                row_cells[4].text = str(member.get("cpu", "N/A"))
                row_cells[5].text = str(member.get("memory", "N/A"))
                row_cells[6].text = member.get("status", "N/A")
        
        doc.add_paragraph(f"\n报告生成时间：{report_time}")

    doc.save(out_path)
    print(f"全新 Word 已生成（每页一台）：{out_path}")


# ------------------- 模板替换（每页插入模板） -------------------

def create_member_table(doc, members_data):
    """根据成员数据创建一个新的表格"""
    if not members_data:
        return None
    
    headers = ["ID/Slot", "角色", "型号", "序列号", "CPU", "内存", "状态"]
    table = doc.add_table(rows=1, cols=len(headers), style="Table Grid")
    table.autofit = True
    
    # 设置表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
    
    # 填充数据
    for member in members_data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(member.get("id", "N/A"))
        row_cells[1].text = member.get("role", "N/A")
        row_cells[2].text = member.get("model", "N/A")
        row_cells[3].text = member.get("sn", "N/A")
        row_cells[4].text = str(member.get("cpu", "N/A"))
        row_cells[5].text = str(member.get("memory", "N/A"))
        row_cells[6].text = member.get("status", "N/A")
    return table._tbl


def replace_template_multi(switches, tmpl_path: Path, out_path: Path):
    doc = Document(tmpl_path)
    set_default_font(doc)

    # 我们需要一个模板的副本用于每次迭代
    # 获取模板的所有顶级元素（段落和表格）
    template_body_elements = [deepcopy(element) for element in doc.element.body]

    # 清空原始文档，稍后逐个添加处理后的页面
    for element in list(doc.element.body):
        doc.element.body.remove(element)
    
    report_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    for i, data in enumerate(switches):
        if i > 0:
            add_page_break(doc)
        
        # 复制模板内容到当前页
        for element in template_body_elements:
            doc.element.body.append(deepcopy(element))

        # 构造占位符映射（大写）
        placeholder_map = {k.upper(): v for k, v in data.items() if not k.startswith("_") and not isinstance(v, (list, dict))}
        placeholder_map["IP"] = data.get("_filename", "N/A")
        placeholder_map["REPORT_TIME"] = report_time
        # 为了兼容性，顶层SN等可以用主设备的信息填充
        placeholder_map["SN"] = data.get("sn", "N/A")
        placeholder_map["CPU_UTILIZATION"] = data.get("cpu_utilization", "N/A")
        placeholder_map["MEMORY_UTILIZATION"] = data.get("memory_utilization", "N/A")
        placeholder_map["MODEL"] = data.get("model", "N/A")
        placeholder_map["IOS_VERSION"] = data.get("ios_version", "N/A")


        # 全局替换简单占位符
        replace_text_in_doc(doc, placeholder_map)

        # 特殊处理 {MEMBER_TABLE}
        for para in doc.paragraphs:
            if '{MEMBER_TABLE}' in para.text:
                para.text = "" # 清空占位符
                # 在这个段落的位置插入新表格
                new_table_xml = create_member_table(doc, data.get("members", []))
                if new_table_xml is not None:
                    p_element = para._p
                    p_element.addnext(new_table_xml)

    doc.save(out_path)
    print(f"模板替换完成（每页一台）：{out_path}")

# ------------------- 主程序 -------------------
def main():
    parser = argparse.ArgumentParser(
        description="""批量读取 *.txt（IP 命名），按 IP 顺序生成 1 个 Word 文档。
支持 Cisco, Huawei, H3C，并能处理堆叠设备。""",
        formatter_class=argparse.RawDescriptionHelpFormatter
    )
    parser.add_argument('-i', '--input', default='.', help='存放 TXT 文件的文件夹，默认当前目录')
    parser.add_argument('-t', '--template', help='Word 模板文件路径（可选，若提供则使用模板替换模式）')
    parser.add_argument('-o', '--output', required=True, help='输出 Word 文件路径')

    args = parser.parse_args()

    txt_dir = Path(args.input)
    if not txt_dir.exists() or not txt_dir.is_dir():
        print(f"输入目录不存在或不是目录：{txt_dir}")
        return

    pattern = re.compile(r"^\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\.txt$")
    txt_files = [p for p in txt_dir.iterdir() if p.is_file() and pattern.match(p.name)]
    if not txt_files:
        print(f"未找到任何 *.txt 文件（IP 命名格式）")
        return

    txt_files = sort_by_ip(txt_files)
    print(f"找到 {len(txt_files)} 个文件，按 IP 排序：")
    for f in txt_files:
        print(f"  → {f.name}")

    switches = [parse_device_info(txt_file) for txt_file in txt_files]

    out_path = Path(args.output)
    if args.template:
        tmpl_path = Path(args.template)
        if not tmpl_path.exists():
            print(f"模板不存在：{tmpl_path}")
            return
        replace_template_multi(switches, tmpl_path, out_path)
    else:
        generate_multi_word(switches, out_path)

if __name__ == "__main__":
    main()
