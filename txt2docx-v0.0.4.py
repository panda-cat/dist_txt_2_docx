# -*- coding: utf-8 -*-
"""
批量读取 *.txt 文件（以IP地址开头命名），按 IP 顺序生成 1 个 Word 文档。
每台交换机一页，支持：
1. 全新表格报告（无模板时）
2. 模板占位符替换（提供模板时）

[!] 核心功能:
- 支持灵活的文件命名（例如 `192.168.1.1.txt`, `192.168.1.1-backup.txt`）。
- 自动识别并解析 Cisco (IOS/XE), Huawei (VRP), H3C (Comware) 设备。
- 适配 Cisco StackWise, Huawei iStack, H3C IRF 堆叠/虚拟化技术。
- 采用优先级策略精确解析主机名，有效避免 LLDP 等信息干扰。
- 支持在 Word 模板的任何位置（包括正文、表格、页眉、页脚）使用任何占位符。
- 使用“增量替换”策略，完美处理模板中包含复杂表格和多个相同占位符的场景。
- 支持在 Word 模板中使用特殊占位符 {MEMBER_TABLE} 动态生成堆叠成员表格。

命令行参数：
  -i, --input     存放 TXT 文件的文件夹，默认当前目录
  -t, --template  Word 模板文件路径（可选）
  -o, --output    输出 Word 文件路径（必填）
"""
import re
import argparse
import ipaddress
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from copy import deepcopy

# ------------------- Word 文档操作辅助函数 -------------------

def add_page_break(doc):
    """添加分页符"""
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def set_default_font(doc, font_name='微软雅黑'):
    """设置文档默认中英文字体"""
    style = doc.styles['Normal']
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def _replace_in_content_part(content_part, placeholder_map):
    """辅助函数，在任何包含 .paragraphs 和 .tables 的对象中递归执行替换。"""
    for para in content_part.paragraphs:
        for run in para.runs:
            for ph, val in placeholder_map.items():
                run.text = run.text.replace(f"{{{ph}}}", str(val))
    
    for table in content_part.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_content_part(cell, placeholder_map)

def replace_placeholders_in_document(doc, placeholder_map):
    """在整个文档的所有部分（包括正文、页眉、页脚）中替换占位符。"""
    # 1. 替换正文内容
    _replace_in_content_part(doc, placeholder_map)
    # 2. 遍历所有“节”，并替换其中的页眉和页脚内容
    for section in doc.sections:
        if section.header: _replace_in_content_part(section.header, placeholder_map)
        if section.footer: _replace_in_content_part(section.footer, placeholder_map)
        if section.first_page_header: _replace_in_content_part(section.first_page_header, placeholder_map)
        if section.first_page_footer: _replace_in_content_part(section.first_page_footer, placeholder_map)

# ------------------- 核心解析逻辑：各厂商解析器 -------------------

def parse_cisco(content: str) -> dict:
    data = {"vendor": "Cisco", "members": []}
    hostname = None
    m_prompt = re.search(r"(\S+?)#", content, re.I)
    if m_prompt: hostname = m_prompt.group(1).strip()
    if not hostname:
        m_config = re.search(r"^\s*hostname\s+(.+?)\s*$", content, re.I | re.M)
        if m_config: hostname = m_config.group(1).strip()
    data["hostname"] = hostname or "N/A"
    
    m = re.search(r"uptime is (.+)", content, re.I); data["uptime"] = m.group(1).strip() if m else "N/A"
    m = re.search(r"Clock is (.+)", content, re.I); data["ntp_status"] = m.group(1).strip() if m else "N/A"
    m_cpu = re.search(r"CPU utilization for five seconds: (\S+)", content, re.I); data["cpu_utilization"] = m_cpu.group(1).split('/')[0] if m_cpu else "N/A"
    m_total = re.search(r"Processor Pool Total:\s+(\d+)", content, re.I)
    m_used = re.search(r"Used:\s+(\d+)", content, re.I)
    if m_total and m_used: data["memory_utilization"] = f"{(int(m_used.group(1)) / int(m_total.group(1)) * 100):.2f}%" if int(m_total.group(1)) > 0 else "0%"
    else: data["memory_utilization"] = "N/A"

    if 'show switch' in content.lower():
        member_matches = re.finditer(r"^\s*([*\d])\s+\S+\s+([\w-]+)\s+([0-9a-f:.]+)\s+(\w+)", content, re.M | re.I)
        members = [{"id": m.group(1).replace('*','').strip(), "model": m.group(2), "mac_address": m.group(3), "status": m.group(4), "sn": "N/A"} for m in member_matches]
        sn_matches = re.finditer(r"Switch\s+(\d+)\s+SERIAL NUMBER\s+:\s+(\S+)", content, re.I | re.M)
        sn_map = {m.group(1): m.group(2) for m in sn_matches}
        for member in members: member["sn"] = sn_map.get(member["id"], "N/A")
        if members: data["members"], data["is_stack"] = members, len(members) > 1
    
    if not data["members"]:
        member = {"id": "1", "status": "Ready", "cpu": data.get("cpu_utilization"), "memory": data.get("memory_utilization")}
        m = re.search(r"System Serial Number\s+:\s+(\S+)", content, re.I); member["sn"] = m.group(1) if m else "N/A"
        m = re.search(r"Model Number\s+:\s+(\S+)", content, re.I); member["model"] = m.group(1) if m else "N/A"
        m = re.search(r"Cisco IOS.*, Version (\S+),", content, re.I); data["ios_version"] = m.group(1) if m else "N/A"
        data["sn"], data["model"] = member["sn"], member["model"]
        data["members"].append(member)
        data["is_stack"] = False
    return data

def parse_huawei(content: str) -> dict:
    data = {"vendor": "Huawei", "members": []}
    hostname = None
    m_prompt = re.search(r"<(\S+?)>", content, re.I)
    if m_prompt: hostname = m_prompt.group(1).strip()
    if not hostname:
        m_config = re.search(r"^\s*(?:sysname|hostname)\s+(.+?)\s*$", content, re.I | re.M)
        if m_config: hostname = m_config.group(1).strip()
    data["hostname"] = hostname or "N/A"

    m = re.search(r"uptime is (.+)", content, re.I); data["uptime"] = m.group(1).strip() if m else "N/A"
    m = re.search(r"clock status\s*:\s*(.+)", content, re.I); data["ntp_status"] = m.group(1).strip() if m else "N/A"
    m = re.search(r"Version \d\.\d+ \((.+?)\)", content, re.I); data["ios_version"] = m.group(1) if m else "N/A"
    
    if 'display device' in content.lower():
        member_matches = re.finditer(r"^\s*(\d+)\s+(\w+)\s+\w+\s+([\w-]+)\s+([0-9A-Z]+)", content, re.M | re.I)
        members = [{"id": m.group(1), "role": m.group(2), "model": m.group(3), "sn": m.group(4), "cpu": "N/A", "memory": "N/A"} for m in member_matches]
        cpu_map = {m.group(1): m.group(2) for m in re.finditer(r"CPU Usage for Slot\s+(\d+)\s+is\s+(\S+)", content, re.I | re.M)}
        mem_map = {m.group(1): m.group(2) for m in re.finditer(r"Memory usage of slot\s+(\d+):\s+(\S+)", content, re.I | re.M)}
        for member in members: member["cpu"], member["memory"] = cpu_map.get(member["id"]), mem_map.get(member["id"])
        if members:
            data["members"], data["is_stack"] = members, len(members) > 1
            for m in members:
                if m.get('role', '').lower() == 'master':
                    data.update({k: m.get(k) for k in ['cpu_utilization', 'memory_utilization', 'sn', 'model']})
                    break
    
    if not data["members"]:
        member = {"id": "1"}
        m = re.search(r"BARCODE\s+:\s+(\S+)", content, re.I | re.M); member['sn'] = m.group(1) if m else "N/A"
        m = re.search(r"ITEM\s+:\s+(\S+)", content, re.I | re.M); member['model'] = m.group(1) if m else "N/A"
        m_cpu = re.search(r"Control Plane\s+CPU Usage is\s*(\S+)", content); member['cpu'] = m_cpu.group(1) if m_cpu else 'N/A'
        m_mem = re.search(r"Memory Using Percentage Is\s*(\S+)", content); member['memory'] = m_mem.group(1) if m_mem else 'N/A'
        data.update({'sn': member['sn'], 'model': member['model'], 'cpu_utilization': member['cpu'], 'memory_utilization': member['memory']})
        data["members"].append(member)
        data["is_stack"] = False
    return data

def parse_h3c(content: str) -> dict:
    data = {"vendor": "H3C", "members": []}
    hostname = None
    m_prompt = re.search(r"<(\S+?)>", content, re.I)
    if m_prompt: hostname = m_prompt.group(1).strip()
    if not hostname:
        m_config = re.search(r"^\s*(?:sysname|hostname)\s+(.+?)\s*$", content, re.I | re.M)
        if m_config: hostname = m_config.group(1).strip()
    data["hostname"] = hostname or "N/A"
    
    m = re.search(r"uptime is (.+)", content, re.I); data["uptime"] = m.group(1).strip() if m else "N/A"
    m = re.search(r"Clock status: (.+)", content, re.I); data["ntp_status"] = m.group(1).strip() if m else "N/A"
    m = re.search(r"Comware Software, Version ([\d\.]+)", content, re.I); data["ios_version"] = m.group(1) if m else "N/A"
    
    if 'display irf' in content.lower() or 'display device' in content.lower():
        member_matches = re.finditer(r"^\s*(\d+)\s+(\w+)\s+([\w-]+)\s+([A-Z0-9]+)", content, re.M | re.I)
        members = [{"id": m.group(1), "role": m.group(2), "model": m.group(3), "sn": m.group(4), "cpu": "N/A", "memory": "N/A"} for m in member_matches]
        cpu_map = {m.group(1): m.group(2) for m in re.finditer(r"Slot\s+(\d+)\s+CPU usage:\s+(\S+)", content, re.I | re.M)}
        mem_map = {m.group(1): m.group(2) for m in re.finditer(r"Slot\s+(\d+)\s+memory usage\s+\(Ratio\):\s+(\S+)", content, re.I | re.M)}
        for member in members: member["cpu"], member["memory"] = cpu_map.get(member["id"]), mem_map.get(member["id"])
        if members:
            data["members"], data["is_stack"] = members, len(members) > 1
            for m in members:
                if m.get('role', '').lower() == 'master':
                    data.update({k: m.get(k) for k in ['cpu_utilization', 'memory_utilization', 'sn', 'model']})
                    break
    
    if not data["members"]:
        member = {"id": "1"}
        m = re.search(r"Device serial number:\s*(\S+)", content, re.I); member['sn'] = m.group(1) if m else "N/A"
        m = re.search(r"Device model:\s*(\S+)", content, re.I); member['model'] = m.group(1) if m else "N/A"
        m_cpu = re.search(r"CPU average usage:\s*(\S+)", content, re.I); member['cpu'] = m_cpu.group(1) if m_cpu else "N/A"
        m_mem = re.search(r"Memory usage:\s*(\S+)", content, re.I); member['memory'] = m_mem.group(1) if m_mem else "N/A"
        data.update({'sn': member['sn'], 'model': member['model'], 'cpu_utilization': member['cpu'], 'memory_utilization': member['memory']})
        data["members"].append(member)
        data["is_stack"] = False
    return data

# ------------------- “控制中心”：自动检测并分发任务 -------------------

def parse_device_info(txt_path: Path, ip_address: str) -> dict:
    with txt_path.open(encoding="utf-8", errors="ignore") as f:
        content = f.read()
    data = {"_filename": ip_address}

    if re.search(r"Cisco IOS|\s#show", content, re.I): parsed_data = parse_cisco(content)
    elif re.search(r"VRP \(R\) software|HUAWEI|<HUAWEI>|display device", content, re.I): parsed_data = parse_huawei(content)
    elif re.search(r"Comware Software|<H3C>|display irf", content, re.I): parsed_data = parse_h3c(content)
    else:
        parsed_data = {"vendor": "Unknown", "is_stack": False, "members": [{}]}
        for m in re.finditer(r"^\s*(.+?)\s*:\s*(.+?)\s*$", content, re.M):
            k, v = m.groups()
            parsed_data[k.lower().replace(" ", "_")] = v.strip()
    data.update(parsed_data)
    return data

# ------------------- 工具函数：按 IP 排序 -------------------

def sort_by_ip(file_list, ip_pattern):
    def ip_key(p):
        match = ip_pattern.match(p.name)
        ip_str = match.group(1) if match else "255.255.255.255"
        try: return ipaddress.ip_address(ip_str)
        except ValueError: return ipaddress.ip_address("255.255.255.255")
    return sorted(file_list, key=ip_key)

# ------------------- 报告生成模块 -------------------

def generate_multi_word(switches, out_path: Path):
    doc = Document(); set_default_font(doc)
    report_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    for i, data in enumerate(switches):
        if i > 0: add_page_break(doc)
        doc.add_heading(f"交换机状态报告 - {data.get('_filename')} ({data.get('hostname')})", level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        for title, headers, keys in [
            ("设备概览", ['主机名', '厂商', '主设备型号', '软件版本'], ['hostname', 'vendor', 'model', 'ios_version']),
            ("运行状态", ['运行时间', 'NTP状态', 'CPU使用率(主)', '内存使用率(主)'], ['uptime', 'ntp_status', 'cpu_utilization', 'memory_utilization'])
        ]:
            doc.add_heading(title, level=2)
            table = doc.add_table(rows=2, cols=len(headers), style="Table Grid")
            for i, h in enumerate(headers): table.cell(0, i).text = h
            for i, k in enumerate(keys): table.cell(1, i).text = str(data.get(k, "N/A"))
        
        doc.add_heading("成员设备详情", level=2)
        members = data.get("members", [])
        if members:
            member_headers = ["ID/Slot", "角色", "型号", "序列号", "CPU", "内存", "状态"]
            table = doc.add_table(rows=1, cols=len(member_headers), style="Table Grid")
            for i, h in enumerate(member_headers): table.rows[0].cells[i].text = h
            for member in members:
                row = table.add_row().cells
                for i, k in enumerate(["id", "role", "model", "sn", "cpu", "memory", "status"]):
                    row[i].text = str(member.get(k, "N/A"))
        doc.add_paragraph(f"\n报告生成时间：{report_time}")
    doc.save(out_path); print(f"全新 Word 报告已生成：{out_path}")

def create_member_table_xml(doc, members_data):
    if not members_data: return None
    headers = ["ID/Slot", "角色", "型号", "序列号", "CPU", "内存", "状态"]
    table = doc.add_table(rows=1, cols=len(headers), style="Table Grid")
    for i, h in enumerate(headers): table.rows[0].cells[i].text = h
    for member in members_data:
        row = table.add_row().cells
        for i, k in enumerate(["id", "role", "model", "sn", "cpu", "memory", "status"]):
            row[i].text = str(member.get(k, "N/A"))
    return table._tbl

def replace_template_multi(switches, tmpl_path: Path, out_path: Path):
    doc = Document(tmpl_path); set_default_font(doc)
    template_body_elements = [deepcopy(element) for element in doc.element.body]
    for element in list(doc.element.body): doc.element.body.remove(element)

    for i, data in enumerate(switches):
        if i > 0: add_page_break(doc)
        paras_before, tables_before = len(doc.paragraphs), len(doc.tables)
        for element in template_body_elements: doc.element.body.append(deepcopy(element))

        placeholder_map = {k.upper(): str(v if v is not None else "N/A") for k, v in data.items() if not k.startswith("_") and not isinstance(v, (list, dict))}
        placeholder_map["IP"] = str(data.get("_filename", "N/A"))
        
        for para in doc.paragraphs[paras_before:]:
            for run in para.runs:
                for ph, val in placeholder_map.items(): run.text = run.text.replace(f"{{{ph}}}", str(val))
        for table in doc.tables[tables_before:]:
            _replace_in_content_part(table, placeholder_map)

        for para in doc.paragraphs[paras_before:]:
            if '{MEMBER_TABLE}' in para.text:
                para.text = "" 
                new_table_xml = create_member_table_xml(doc, data.get("members", []))
                if new_table_xml:
                    p_element = para._p
                    p_element.addnext(new_table_xml)
                    parent = new_table_xml.getparent()
                    if parent is not None: parent.remove(new_table_xml)
    
    final_ph_map = {"REPORT_TIME": datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
    if switches:
        last_data = switches[-1]
        final_ph_map.update({k.upper(): str(v if v is not None else "N/A") for k, v in last_data.items() if not k.startswith("_") and not isinstance(v, (list, dict))})
        final_ph_map["IP"] = str(last_data.get("_filename", "N/A"))
    replace_placeholders_in_document(doc, final_ph_map)
    doc.save(out_path); print(f"模板替换完成：{out_path}")

# ------------------- 主程序入口 -------------------

def main():
    parser = argparse.ArgumentParser(description="批量读取TXT日志，生成多厂商、支持堆叠的Word巡检报告。", formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument('-i', '--input', default='.', help='存放 TXT 文件的文件夹')
    parser.add_argument('-t', '--template', help='Word 模板文件路径（可选）')
    parser.add_argument('-o', '--output', required=True, help='输出 Word 文件路径')
    args = parser.parse_args()

    txt_dir = Path(args.input)
    if not txt_dir.is_dir(): print(f"错误：输入路径不是一个有效的目录：{txt_dir}"); return
    
    ip_pattern = re.compile(r"^(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}).*\.txt$")
    txt_files = [p for p in txt_dir.iterdir() if p.is_file() and ip_pattern.match(p.name)]
    if not txt_files: print(f"在目录 '{txt_dir}' 中未找到任何以IP地址开头的 .txt 文件。"); return

    txt_files = sort_by_ip(txt_files, ip_pattern)
    print(f"找到 {len(txt_files)} 个文件，已按 IP 地址排序：")
    for f in txt_files: print(f"  → {f.name}")

    switches = [parse_device_info(p, ip_pattern.match(p.name).group(1)) for p in txt_files]

    if args.template:
        tmpl_path = Path(args.template)
        if not tmpl_path.is_file(): print(f"错误：模板文件不存在：{tmpl_path}"); return
        replace_template_multi(switches, tmpl_path, Path(args.output))
    else:
        generate_multi_word(switches, Path(args.output))

if __name__ == "__main__":
    main()
